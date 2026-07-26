"""
Generate documentation .docx files for Website Desa Makmurjaya.
1. Panduan Pengelola CMS (for village admin)
2. Architecture & Extension Guide (for future developers)

Run: python docs/generate_docs.py
Output: docs/Panduan-Pengelola-CMS.docx, docs/Architecture-Extension-Guide.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_image_placeholder(doc, description):
    """Add a bordered placeholder paragraph indicating where a screenshot goes."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[SCREENSHOT: {description}]")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    # Add border-like spacing
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_step(doc, number, text):
    """Add a numbered step."""
    p = doc.add_paragraph()
    run = p.add_run(f"Langkah {number}: ")
    run.bold = True
    p.add_run(text)
    return p


def generate_panduan_cms():
    """Generate the CMS guide for village administrators."""
    doc = Document()

    # Title page
    doc.add_paragraph()
    title = doc.add_heading("PANDUAN PENGELOLA WEBSITE", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Desa Makmurjaya", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Panduan lengkap penggunaan Content Management System (CMS)\n"
                          "untuk mengelola konten website desa.\n\n"
                          "Ditujukan untuk: Pak Hendra (Pengelola Website Desa)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph("Disusun oleh: Muhamad Tristan Farand\nKKN IPB 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading("Daftar Isi", level=1)
    toc_items = [
        "1. Pengenalan Sistem",
        "2. Cara Login ke CMS",
        "3. Mengelola Berita",
        "4. Mengelola Pengumuman & Agenda",
        "5. Mengelola UMKM (Belanja)",
        "6. Mengelola Wisata",
        "7. Mengelola Galeri Foto",
        "8. Mengelola Perangkat Desa",
        "9. Mengelola FAQ",
        "10. Mengelola Infografis",
        "11. Pengaturan Situs",
        "12. Troubleshooting (Masalah Umum)",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # Chapter 1: Introduction
    doc.add_heading("1. Pengenalan Sistem", level=1)
    doc.add_paragraph(
        "Website Desa Makmurjaya dikelola menggunakan sistem CMS (Content Management System) "
        "yang memungkinkan Anda mengedit konten website langsung dari browser — "
        "seperti mengedit dokumen Word, tetapi hasilnya langsung tampil di website."
    )
    doc.add_heading("Bagaimana Cara Kerjanya?", level=2)
    doc.add_paragraph(
        "1. Anda login ke halaman admin CMS melalui browser\n"
        "2. Edit konten yang diinginkan (berita, pengumuman, foto, dll)\n"
        "3. Klik tombol 'Publish' atau 'Simpan'\n"
        "4. Tunggu 1-2 menit — website akan otomatis berubah"
    )
    create_image_placeholder(doc, "Diagram alur: Admin edit → Simpan → Website berubah otomatis (1-2 menit)")
    doc.add_heading("Alamat Akses CMS", level=2)
    doc.add_paragraph("Buka browser (Chrome/Firefox) dan ketik alamat berikut:")
    p = doc.add_paragraph()
    run = p.add_run("https://jeff146354.github.io/website-desa-makmurjaya/admin")
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(
        "Catatan: Anda memerlukan akun GitHub yang sudah diberikan akses oleh pengembang."
    )
    doc.add_page_break()

    # Chapter 2: Login
    doc.add_heading("2. Cara Login ke CMS", level=1)
    add_step(doc, 1, "Buka browser dan akses alamat CMS di atas.")
    create_image_placeholder(doc, "Tampilan halaman login CMS — tombol 'Login with GitHub'")
    add_step(doc, 2, 'Klik tombol "Login with GitHub".')
    add_step(doc, 3, "Masukkan username dan password akun GitHub Anda.")
    create_image_placeholder(doc, "Form login GitHub — username & password fields")
    add_step(doc, 4, "Jika diminta, klik 'Authorize' untuk memberikan izin akses.")
    add_step(doc, 5, "Anda akan masuk ke Dashboard CMS. Selamat!")
    create_image_placeholder(doc, "Dashboard CMS setelah login — menampilkan daftar collection (Berita, Pengumuman, dll)")
    doc.add_page_break()

    # Chapter 3: Berita
    doc.add_heading("3. Mengelola Berita", level=1)
    doc.add_heading("Menambah Berita Baru", level=2)
    add_step(doc, 1, 'Di dashboard CMS, klik menu "Berita" di sidebar kiri.')
    add_step(doc, 2, 'Klik tombol "New Artikel Berita" di pojok kanan atas.')
    create_image_placeholder(doc, "Halaman daftar berita dengan tombol 'New Artikel Berita' yang ditandai")
    add_step(doc, 3, "Isi form berita:")
    doc.add_paragraph(
        "• Judul — Judul berita yang menarik\n"
        "• Slug URL — Alamat halaman (otomatis dari judul, huruf kecil + tanda hubung)\n"
        "• Tanggal Publikasi — Pilih tanggal\n"
        "• Foto Thumbnail — Upload foto (maks 5MB, format JPG/PNG/WebP)\n"
        "• Cuplikan — Ringkasan singkat (maks 300 karakter)\n"
        "• Isi Berita — Tulis isi berita lengkap (mendukung format teks tebal, miring, gambar)"
    )
    create_image_placeholder(doc, "Form editor berita — menunjukkan field Judul, Tanggal, Foto, dan editor Markdown")
    add_step(doc, 4, 'Klik tombol "Publish" di pojok kanan atas.')
    doc.add_paragraph("Tips: Ukuran foto yang ideal adalah lebar 800-1200px dengan format WebP atau JPG.")
    doc.add_paragraph(
        "Catatan: Panduan ini ditujukan untuk Pak Hendra selaku pengelola website desa. "
        "Jika ada pertanyaan, silakan hubungi pengembang."
    )
    doc.add_heading("Mengedit Berita", level=2)
    doc.add_paragraph("Klik judul berita yang ingin diedit dari daftar → ubah isinya → klik Publish.")
    doc.add_heading("Menghapus Berita", level=2)
    doc.add_paragraph('Buka berita → klik tombol "Delete" di bagian bawah → konfirmasi.')
    doc.add_page_break()

    # Chapter 4: Pengumuman
    doc.add_heading("4. Mengelola Pengumuman & Agenda", level=1)
    doc.add_paragraph("Pengumuman dan Agenda dikelola di menu yang sama, dibedakan oleh field 'Kategori'.")
    doc.add_heading("Perbedaan Pengumuman vs Agenda", level=2)
    doc.add_paragraph(
        "• Pengumuman — informasi umum dari desa (contoh: libur pelayanan)\n"
        "• Agenda — kegiatan yang memiliki tanggal & waktu pelaksanaan (contoh: musyawarah desa)"
    )
    doc.add_heading("Menambah Pengumuman/Agenda", level=2)
    add_step(doc, 1, 'Klik menu "Pengumuman" → "New Pengumuman".')
    add_step(doc, 2, "Isi field: Judul, Slug URL, Kategori (pilih Pengumuman atau Agenda), Tanggal.")
    add_step(doc, 3, "Jika Agenda: isi juga Tanggal Pelaksanaan dan Waktu Pelaksanaan.")
    add_step(doc, 4, "Opsional: lampirkan file (poster, surat undangan) di bagian Lampiran.")
    add_step(doc, 5, 'Tulis isi pengumuman di editor → Publish.')
    create_image_placeholder(doc, "Form pengumuman dengan field Kategori (dropdown Pengumuman/Agenda) yang ditandai")
    doc.add_page_break()

    # Chapter 5: UMKM
    doc.add_heading("5. Mengelola UMKM (Belanja)", level=1)
    add_step(doc, 1, 'Klik menu "UMKM" → "New UMKM".')
    add_step(doc, 2, "Isi informasi usaha:")
    doc.add_paragraph(
        "• Nama Usaha — Nama lengkap usaha/toko\n"
        "• Slug URL — Alamat halaman\n"
        "• Pemilik — Nama pemilik usaha\n"
        "• Kategori — Jenis usaha (Kuliner, Kerajinan, Hiburan, dll)\n"
        "• Jenis Produk — Produk utama yang dijual\n"
        "• Range Harga — Rentang harga (contoh: Rp 5.000 - Rp 50.000)\n"
        "• Nomor WhatsApp — Format: 628XXXXXXXXX (tanpa spasi/tanda hubung)\n"
        "• Foto Utama — Foto produk/toko (maks 5MB)\n"
        "• Galeri Produk — Foto-foto tambahan\n"
        "• Deskripsi Usaha — Deskripsi lengkap di editor"
    )
    create_image_placeholder(doc, "Form UMKM — field Nama Usaha, Kategori, WhatsApp, dan upload Foto")
    add_step(doc, 3, 'Klik "Publish".')
    doc.add_paragraph(
        "PENTING: Format nomor WhatsApp harus diawali 628 (bukan 08). "
        "Contoh: 0858-8226-4441 → tulis 6285882264441"
    )
    doc.add_page_break()

    # Chapter 6: Wisata
    doc.add_heading("6. Mengelola Wisata", level=1)
    add_step(doc, 1, 'Klik menu "Destinasi Wisata" → "New Destinasi Wisata".')
    add_step(doc, 2, "Isi: Nama Destinasi, Deskripsi Singkat, Jam Operasional, Harga Tiket, Kontak, Petunjuk Akses.")
    add_step(doc, 3, "Opsional: isi koordinat peta (Latitude & Longitude) agar peta tampil di halaman wisata.")
    add_step(doc, 4, "Upload Foto Utama dan Galeri Foto.")
    add_step(doc, 5, "Tulis deskripsi lengkap di editor → Publish.")
    create_image_placeholder(doc, "Form wisata — field koordinat peta, galeri foto, dan deskripsi")
    doc.add_paragraph("Tips: Untuk mendapatkan koordinat, buka Google Maps → klik kanan lokasi → salin koordinat.")
    doc.add_page_break()

    # Chapter 7: Galeri
    doc.add_heading("7. Mengelola Galeri Foto", level=1)
    add_step(doc, 1, 'Klik menu "Galeri Foto" → "New Galeri Foto".')
    add_step(doc, 2, "Isi: Judul Foto, Upload File Foto, Caption (keterangan), Album/Kategori.")
    add_step(doc, 3, "Atur Urutan (angka kecil tampil lebih dulu).")
    add_step(doc, 4, "Publish.")
    create_image_placeholder(doc, "Form galeri — upload foto, isi caption, pilih album")
    doc.add_page_break()

    # Chapter 8: Perangkat Desa
    doc.add_heading("8. Mengelola Perangkat Desa", level=1)
    add_step(doc, 1, 'Klik menu "Perangkat Desa" → "New Perangkat Desa".')
    add_step(doc, 2, "Isi: Nama, Jabatan, Upload Foto, Urutan Tampil, Level Hierarki.")
    doc.add_paragraph(
        "Level Hierarki:\n"
        "• 1 = Kepala Desa\n"
        "• 2 = Sekretaris Desa\n"
        "• 3 = Kepala Seksi/Kaur\n"
        "• 4 = Staf"
    )
    add_step(doc, 3, "Isi Tugas Pokok di editor markdown → Publish.")
    create_image_placeholder(doc, "Form perangkat desa — field hierarki, urutan, foto")
    doc.add_page_break()

    # Chapter 9: FAQ
    doc.add_heading("9. Mengelola FAQ", level=1)
    add_step(doc, 1, 'Klik menu "FAQ (Tanya Jawab)" → "New FAQ".')
    add_step(doc, 2, "Isi: Pertanyaan, Kategori (opsional), Urutan.")
    add_step(doc, 3, "Tulis jawaban di editor markdown → Publish.")
    doc.add_page_break()

    # Chapter 10: Infografis
    doc.add_heading("10. Mengelola Infografis", level=1)
    doc.add_paragraph(
        "Data infografis (Penduduk, APBDes, IDM, Stunting, Bantuan Sosial, SDGs) "
        "dikelola melalui menu 'Infografis' di CMS."
    )
    doc.add_heading("Update Data Penduduk", level=2)
    add_step(doc, 1, 'Klik "Infografis" → "Data Penduduk".')
    add_step(doc, 2, "Update angka: Total Penduduk, KK, Laki-laki, Perempuan.")
    add_step(doc, 3, "Tambah/edit data kelompok umur, per dusun, per agama, dll.")
    add_step(doc, 4, "Save.")
    doc.add_heading("Update Data IDM", level=2)
    add_step(doc, 1, 'Klik "Infografis" → "Data IDM".')
    add_step(doc, 2, "Update Skor IDM, Status, IKS/IKE/IKL.")
    add_step(doc, 3, "Untuk tabel indikator: edit skor masing-masing indikator di list 'Tabel Indikator'.")
    create_image_placeholder(doc, "Halaman edit Infografis IDM — menunjukkan field skor dan tabel indikator")
    doc.add_heading("Update Data APBDes", level=2)
    doc.add_paragraph("Klik 'Data APBDes' → isi tahun anggaran, pendapatan, belanja, rincian → Save.")
    doc.add_page_break()

    # Chapter 11: Pengaturan Situs
    doc.add_heading("11. Pengaturan Situs", level=1)
    doc.add_paragraph(
        "Menu ini mengatur informasi global yang tampil di seluruh halaman website "
        "(header, footer, homepage hero)."
    )
    add_step(doc, 1, 'Klik "Pengaturan Situs" → "Informasi Umum Desa".')
    add_step(doc, 2, "Edit field yang diinginkan:")
    doc.add_paragraph(
        "• Nama Desa — Tampil di header & footer\n"
        "• Lokasi Singkat — Tampil di bawah nama desa\n"
        "• Tagline Hero — Teks di bagian hero homepage\n"
        "• Gambar Hero — Foto utama homepage\n"
        "• Alamat Lengkap — Tampil di footer & halaman kontak\n"
        "• Jam Layanan — Tampil di footer & halaman kontak\n"
        "• Nama & Foto Kepala Desa — Untuk halaman profil\n"
        "• Teks Copyright — Teks di bagian paling bawah footer"
    )
    add_step(doc, 3, "Save.")
    create_image_placeholder(doc, "Form Pengaturan Situs — semua field global (nama desa, alamat, tagline, dll)")
    doc.add_page_break()

    # Chapter 12: Troubleshooting
    doc.add_heading("12. Troubleshooting (Masalah Umum)", level=1)
    problems = [
        ("Perubahan tidak muncul di website",
         "Tunggu 1-2 menit. Website perlu waktu untuk rebuild setelah Anda menyimpan. "
         "Coba refresh halaman (Ctrl+F5)."),
        ("Tidak bisa login",
         "Pastikan akun GitHub Anda sudah diberikan akses ke repository. "
         "Hubungi pengembang jika masih bermasalah."),
        ("Foto tidak muncul",
         "Pastikan format foto adalah JPG, PNG, atau WebP dan ukuran kurang dari 5MB. "
         "Nama file sebaiknya tanpa spasi atau karakter khusus."),
        ("Salah input data",
         "Buka kembali item yang salah → edit → Publish ulang. "
         "Semua perubahan ter-versioning di GitHub, jadi data lama tidak hilang."),
        ("Website error / tidak bisa diakses",
         "Kemungkinan ada error saat build. Hubungi pengembang untuk mengecek "
         "status build di GitHub Actions."),
    ]
    for problem, solution in problems:
        doc.add_heading(f"❓ {problem}", level=2)
        doc.add_paragraph(f"✅ Solusi: {solution}")

    doc.add_page_break()
    doc.add_heading("Kontak Bantuan Teknis", level=1)
    doc.add_paragraph(
        "Jika mengalami masalah yang tidak bisa diselesaikan dengan panduan di atas, "
        "hubungi pengembang:\n\n"
        "Nama: Muhamad Tristan Farand\n"
        "Program: KKN IPB 2026\n"
        "GitHub: Jeff146354\n\n"
        "Pengelola Website Desa: Pak Hendra\n"
        "Kepala Desa: HJ. Nining Nurnaningsih, S.PD"
    )

    doc.save("docs/Panduan-Pengelola-CMS.docx")
    print("✓ Generated: docs/Panduan-Pengelola-CMS.docx")


def generate_architecture_guide():
    """Generate the technical architecture & extension guide."""
    doc = Document()

    # Title
    doc.add_paragraph()
    title = doc.add_heading("ARCHITECTURE & EXTENSION GUIDE", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Website Desa Makmurjaya", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        "Technical documentation for future developers.\n"
        "Created by: Muhamad Tristan Farand · KKN IPB 2026"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. Architecture Overview",
        "2. Technology Stack & Justifications",
        "3. Project Structure",
        "4. Data Flow",
        "5. How to Add a New Page",
        "6. How to Add a New CMS Collection",
        "7. How to Add a New Infografis Tab",
        "8. How to Change the Design System",
        "9. Deployment & CI/CD",
        "10. Known Limitations & Tech Debt",
        "11. Recovery Playbook",
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # Chapter 1
    doc.add_heading("1. Architecture Overview", level=1)
    doc.add_paragraph(
        "This is a static site built with Astro 4.x, styled with Tailwind CSS, "
        "content-managed via Decap CMS (Git-backed), and hosted on GitHub Pages (free). "
        "All content lives as Markdown/JSON files in the same Git repository as source code."
    )
    doc.add_paragraph(
        "Flow: Admin edits in CMS → Git commit → GitHub Actions builds → "
        "Static files deployed to GitHub Pages CDN."
    )
    create_image_placeholder(doc,
        "Architecture diagram: Browser(Admin) → Decap CMS → GitHub Repo → "
        "GitHub Actions (astro build) → GitHub Pages → Browser(Visitor)")
    doc.add_page_break()

    # Chapter 2
    doc.add_heading("2. Technology Stack & Justifications", level=1)
    techs = [
        ("Astro 4.x", "Static site generator. Zero JS by default, Islands Architecture for interactive components. Chosen over Next.js (overkill for static) and Hugo (less flexible templating)."),
        ("Tailwind CSS 3.x", "Utility-first CSS. Custom design tokens in tailwind.config.mjs. Chosen over Bootstrap (too opinionated) and plain CSS (maintenance burden)."),
        ("Decap CMS 3.x", "Browser-based, Git-backed. No server needed. Chosen over WordPress (requires hosting), Strapi (needs backend server), Sanity (paid at scale)."),
        ("GitHub Pages", "Free static hosting with CDN. Chosen over Netlify/Vercel (simpler, no account sprawl)."),
        ("Leaflet.js", "Open-source maps. No API key needed. Chosen over Google Maps (requires billing)."),
        ("Pagefind", "Client-side search built at deploy time. Zero runtime cost. Chosen over Algolia (paid)."),
        ("Vitest + fast-check", "Unit + property-based testing. Fast, ESM-native."),
        ("TypeScript", "Type safety for maintainability. Strict mode enabled."),
    ]
    for name, reason in techs:
        doc.add_heading(name, level=2)
        doc.add_paragraph(reason)
    doc.add_page_break()

    # Chapter 3
    doc.add_heading("3. Project Structure", level=1)
    structure = """
src/
├── components/       → Astro components (by feature)
│   ├── common/       → Header, Footer, Navigation, SEOHead, etc.
│   ├── beranda/      → Homepage sections
│   ├── berita/       → BeritaCard
│   ├── pengumuman/   → PengumumanCard
│   ├── umkm/         → UMKMCard, UMKMFilter
│   ├── wisata/       → WisataCard, WisataDetail, WisataGaleri
│   ├── pemerintahan/ → OrgChart, PerangkatCard
│   ├── galeri/       → GaleriGrid, Lightbox
│   ├── kontak/       → KontakForm, PetaKantor
│   ├── profil/       → ProfilNav
│   └── ui/           → Pagination, Badge, EmptyState, Accordion
├── content/          → Astro Content Collections (Markdown/JSON)
│   ├── berita/       → News articles (.md)
│   ├── pengumuman/   → Announcements (.md)
│   ├── wisata/       → Tourism destinations (.md)
│   ├── umkm/         → UMKM businesses (.md)
│   ├── galeri/       → Photo gallery (.json)
│   ├── perangkat-desa/ → Government officials (.md)
│   ├── faq/          → FAQ entries (.md)
│   ├── halaman-statis/ → Static pages (.md)
│   └── infografis/   → Data JSON (penduduk, idm, apbdes, etc.)
├── data/             → Global site settings
│   └── situs.json    → Site name, contact, hero text (CMS-editable)
├── layouts/          → BaseLayout, ContentLayout
├── pages/            → File-based routing
├── styles/           → global.css (Tailwind + component classes)
└── utils/            → Utility functions (date, currency, pagination, etc.)

public/
├── admin/            → Decap CMS (index.html with inline config)
├── uploads/          → Media files uploaded via CMS
└── robots.txt
    """.strip()
    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.size = Pt(8)
    run.font.name = "Consolas"
    doc.add_page_break()

    # Chapter 4
    doc.add_heading("4. Data Flow", level=1)
    doc.add_heading("Content Collections (Markdown)", level=2)
    doc.add_paragraph(
        "Berita, Pengumuman, Wisata, UMKM, FAQ, Perangkat Desa, and Halaman Statis "
        "use Astro Content Collections defined in src/content/config.ts. "
        "Each collection has a Zod schema for type validation at build time."
    )
    doc.add_heading("Infografis Data (JSON)", level=2)
    doc.add_paragraph(
        "Infografis data lives in src/content/infografis/*.json. "
        "These are imported directly in pages/infografis.astro. "
        "Decap CMS edits these files via the 'files' collection type."
    )
    doc.add_heading("Site Settings (JSON)", level=2)
    doc.add_paragraph(
        "Global site text (nama desa, alamat, tagline) lives in src/data/situs.json. "
        "Imported by Header, Footer, and Homepage components. "
        "Editable via CMS 'Pengaturan Situs' menu."
    )
    doc.add_heading("getPath() Utility", level=2)
    doc.add_paragraph(
        "All internal links use getPath('/path') which prepends the base URL "
        "(/website-desa-makmurjaya/) for GitHub Pages subdirectory deployment. "
        "Never hardcode paths — always use getPath()."
    )
    doc.add_page_break()

    # Chapter 5
    doc.add_heading("5. How to Add a New Page", level=1)
    doc.add_paragraph("Example: Adding a '/layanan' page.")
    add_step(doc, 1, "Create src/pages/layanan.astro")
    add_step(doc, 2, "Import BaseLayout and getPath:")
    doc.add_paragraph(
        '---\n'
        'import BaseLayout from "../layouts/BaseLayout.astro";\n'
        'import { getPath } from "../utils/links";\n'
        '---\n'
        '<BaseLayout title="Layanan" description="...">\n'
        '  <div class="page-header">...</div>\n'
        '  <div class="max-w-7xl mx-auto px-4 md:px-8 py-8">...</div>\n'
        '</BaseLayout>'
    )
    add_step(doc, 3, "Add navigation link in src/components/common/Header.astro (both desktop navlinks and mobile drawer)")
    add_step(doc, 4, "Build and test: npx astro build")
    doc.add_page_break()

    # Chapter 6
    doc.add_heading("6. How to Add a New CMS Collection", level=1)
    add_step(doc, 1, "Define schema in src/content/config.ts using Zod")
    add_step(doc, 2, "Create the folder: src/content/your-collection/")
    add_step(doc, 3, "Add collection to Decap CMS config in public/admin/index.html")
    add_step(doc, 4, "Create the page that queries this collection using getCollection()")
    doc.add_paragraph(
        "Important: The collection name in config.ts must match the folder name "
        "and the Decap CMS collection name exactly."
    )
    doc.add_page_break()

    # Chapter 7
    doc.add_heading("7. How to Add a New Infografis Tab", level=1)
    add_step(doc, 1, "Create src/content/infografis/your-tab.json with initial data")
    add_step(doc, 2, "Add CMS file entry in public/admin/index.html under the 'infografis' collection")
    add_step(doc, 3, "In src/pages/infografis.astro: import the JSON, add a tab button, add a panel div")
    add_step(doc, 4, "The tab switching JS is already generic — it works on aria-controls/role attributes")
    doc.add_page_break()

    # Chapter 8
    doc.add_heading("8. How to Change the Design System", level=1)
    doc.add_paragraph("All design tokens are in tailwind.config.mjs:")
    doc.add_paragraph(
        "• Colors: primary (greens), gold (accents), ink, muted, paper, line\n"
        "• Fonts: font-display (Lexend), font-body (Inter)\n"
        "• Shadows: card, elevated\n\n"
        "Component classes are in src/styles/global.css (@layer components):\n"
        "• .card, .stat-card, .pill, .btn-primary, .btn-outline, .btn-ghost\n"
        "• .page-header, .tab-bar, .tab-item, .empty-state, .page-btn\n"
        "• .progress-wrap, .progress-fill, .table-clean"
    )
    doc.add_paragraph("To change colors: update tailwind.config.mjs → all components update automatically.")
    doc.add_page_break()

    # Chapter 9
    doc.add_heading("9. Deployment & CI/CD", level=1)
    doc.add_paragraph(
        "Deployment is fully automated via GitHub Actions (.github/workflows/deploy.yml):\n\n"
        "1. Push to main branch\n"
        "2. GitHub Actions triggers: checkout → Node 20 → npm ci → astro build → pagefind → upload artifact\n"
        "3. Deploy to GitHub Pages\n\n"
        "Build time: ~30-60 seconds. Available at https://jeff146354.github.io/website-desa-makmurjaya/"
    )
    doc.add_heading("Environment Variables", level=2)
    doc.add_paragraph(
        "• FORMSPREE_ENDPOINT — Set in GitHub repo secrets (Settings → Secrets → Actions)\n"
        "• No other env vars needed for basic operation"
    )
    doc.add_heading("Custom Domain Setup", level=2)
    doc.add_paragraph(
        "1. Add domain to public/CNAME\n"
        "2. Update 'site' in astro.config.mjs\n"
        "3. Update 'base' to '/' if using root domain\n"
        "4. Configure DNS (CNAME record pointing to jeff146354.github.io)"
    )
    doc.add_page_break()

    # Chapter 10
    doc.add_heading("10. Known Limitations & Tech Debt", level=1)
    limitations = [
        "No real-time preview — site rebuilds on each commit (1-2 min delay)",
        "Pagefind search index only updates on deploy, not on CMS save",
        "Image optimization happens at build time only (no runtime CDN resizing)",
        "No image upload size enforcement — CMS only shows a hint (5MB)",
        "Formspree free tier: 50 submissions/month",
        "CMS requires GitHub account — village admin needs a GitHub login",
        "Decap CMS OAuth proxy depends on Netlify (external dependency)",
        "No offline editing capability",
        "No multi-language support (Indonesian only)",
        "Mobile bottom nav overlaps content at bottom — 64px spacer added",
    ]
    for lim in limitations:
        doc.add_paragraph(f"• {lim}")
    doc.add_page_break()

    # Chapter 11
    doc.add_heading("11. Recovery Playbook", level=1)
    doc.add_heading("Build fails after CMS edit", level=2)
    doc.add_paragraph(
        "1. Check GitHub Actions tab for error message\n"
        "2. Most likely: invalid frontmatter in a .md file (missing required field)\n"
        "3. Fix: edit the problematic file via CMS or directly in GitHub\n"
        "4. Push fix → build auto-retries"
    )
    doc.add_heading("Rollback a bad commit", level=2)
    doc.add_paragraph(
        "git log --oneline  # find the good commit hash\n"
        "git revert HEAD    # revert last commit\n"
        "git push origin main"
    )
    doc.add_heading("Set up on a new GitHub account", level=2)
    doc.add_paragraph(
        "1. Fork or transfer the repository\n"
        "2. Enable GitHub Pages (Settings → Pages → Source: GitHub Actions)\n"
        "3. Update repo name in public/admin/index.html (backend.repo field)\n"
        "4. Set up Netlify OAuth proxy (or use github-oauth-app alternative)\n"
        "5. Update astro.config.mjs site & base URLs"
    )
    doc.add_heading("CMS login not working", level=2)
    doc.add_paragraph(
        "1. Verify Netlify site_id in public/admin/index.html matches your Netlify app\n"
        "2. Check GitHub OAuth App callback URL\n"
        "3. Ensure the GitHub user has write access to the repository"
    )

    doc.save("docs/Architecture-Extension-Guide.docx")
    print("✓ Generated: docs/Architecture-Extension-Guide.docx")


if __name__ == "__main__":
    generate_panduan_cms()
    generate_architecture_guide()
    print("\nDone! Both documents generated in docs/ folder.")
